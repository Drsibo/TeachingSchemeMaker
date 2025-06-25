from PySide6.QtCore import QThread, Signal
import os
from docx import Document

class BatchReplaceWorker(QThread):
    log_signal = Signal(str)  # 定义日志信号

    def __init__(self):
        super().__init__()
        self.old_text = None
        self.new_text = None

    def set_replace_text(self, old_text, new_text):
        """设置待替换内容和替换后的内容"""
        self.old_text = old_text
        self.new_text = new_text

    def run(self):
        try:
            # 定义要处理的文件夹
            output_dir = 'output'

            # 检查文件夹是否存在
            if not os.path.exists(output_dir):
                self.log_signal.emit(f"文件夹不存在：{output_dir}")
                return

            # 遍历文件夹中的所有文件
            for filename in os.listdir(output_dir):
                if filename.endswith('.docx'):
                    file_path = os.path.join(output_dir, filename)
                    self.replace_in_word(file_path, self.old_text, self.new_text)

            self.log_signal.emit("批量替换完成！")
        except Exception as e:
            self.log_signal.emit(f"批量替换时出错：{e}")

    def replace_in_word(self, file_path, old_text, new_text):
        """在Word文档中替换指定内容"""
        try:
            # 打开Word文档
            doc = Document(file_path)

            # 遍历文档中的每个段落
            for paragraph in doc.paragraphs:
                self.replace_in_paragraph(paragraph, old_text, new_text)

            # 遍历文档中的每个表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_in_paragraph(paragraph, old_text, new_text)

            # 保存修改后的文档
            doc.save(file_path)
            self.log_signal.emit(f"文件 '{file_path}' 替换完成")
        except Exception as e:
            self.log_signal.emit(f"处理文件 '{file_path}' 时出错：{e}")

    def replace_in_paragraph(self, paragraph, old_text, new_text):
        """在段落中替换指定内容，保持格式"""
        if old_text in paragraph.text:
            # 获取段落的运行列表
            runs = list(paragraph.runs)
            # 清空段落
            paragraph.clear()
            # 创建一个标志，用于判断是否已经找到待替换的文本
            replaced = False
            # 遍历运行列表
            for run in runs:
                text = run.text
                if old_text in text and not replaced:
                    # 分割文本
                    parts = text.split(old_text)
                    # 添加分割后的文本和替换文本
                    for i, part in enumerate(parts):
                        paragraph.add_run(part)
                        if i < len(parts) - 1:
                            # 获取原始运行的格式
                            new_run = paragraph.add_run(new_text)
                            # 复制原始运行的格式到新运行
                            self.copy_run_format(run, new_run)
                            replaced = True
                else:
                    # 直接添加原始运行
                    new_run = paragraph.add_run(text)
                    self.copy_run_format(run, new_run)

    def copy_run_format(self, source_run, target_run):
        """复制运行的格式"""
        # 复制常规属性
        target_run.bold = source_run.bold if hasattr(source_run, 'bold') else None
        target_run.italic = source_run.italic if hasattr(source_run, 'italic') else None
        target_run.underline = source_run.underline if hasattr(source_run, 'underline') else None
        target_run.font.name = source_run.font.name if hasattr(source_run, 'font') and hasattr(source_run.font, 'name') else None
        target_run.font.size = source_run.font.size if hasattr(source_run, 'font') and hasattr(source_run.font, 'size') else None
        target_run.font.color.rgb = source_run.font.color.rgb if hasattr(source_run, 'font') and hasattr(source_run.font, 'color') and hasattr(source_run.font.color, 'rgb') else None
        target_run.font.highlight_color = source_run.font.highlight_color if hasattr(source_run, 'font') and hasattr(source_run.font, 'highlight_color') else None
        target_run.font.all_caps = source_run.font.all_caps if hasattr(source_run, 'font') and hasattr(source_run.font, 'all_caps') else None
        target_run.font.small_caps = source_run.font.small_caps if hasattr(source_run, 'font') and hasattr(source_run.font, 'small_caps') else None
        target_run.font.shadow = source_run.font.shadow if hasattr(source_run, 'font') and hasattr(source_run.font, 'shadow') else None
        target_run.font.hidden = source_run.font.hidden if hasattr(source_run, 'font') and hasattr(source_run.font, 'hidden') else None