from PySide6.QtCore import QThread, Signal
import os
import csv
import shutil
import re
from docx import Document  # 用于读取Word模板文件
from utils.log_util import LogUtil

class CSVProcessWorker(QThread):
    log_signal = Signal(str)  # 定义日志信号

    def __init__(self):
        super().__init__()
        self.log_util = LogUtil(level='INFO')
        self.logger = self.log_util.get_logger()
        self.template_path = 'template.docx'  # 模板文件路径
        self.expected_columns = self.get_placeholder_count()  # 从模板文件中获取占位符数量

    def get_placeholder_count(self):
        """从Word模板文件中获取占位符数量"""
        try:
            if not os.path.exists(self.template_path):
                self.logger.error(f"模板文件不存在：{self.template_path}")
                return 11  # 默认值，如果模板文件不存在

            doc = Document(self.template_path)
            placeholder_set = set()  # 使用集合来存储占位符，自动去除重复项

            # 定义正则表达式来匹配占位符 {{1}}, {{2}}, 等
            placeholder_pattern = re.compile(r'\{\{\s*(\d+)\s*\}\}')

            # 遍历文档中的每个段落
            for paragraph in doc.paragraphs:
                matches = placeholder_pattern.findall(paragraph.text)
                if matches:
                    self.logger.info(f"段落中匹配到的占位符: {matches}")
                placeholder_set.update(matches)

            # 遍历文档中的每个表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = placeholder_pattern.findall(paragraph.text)
                            if matches:
                                self.logger.info(f"表格单元格中匹配到的占位符: {matches}")
                            placeholder_set.update(matches)

            placeholder_count = len(placeholder_set)  # 集合的大小即为占位符的数量
            self.logger.info(f"模板文件中共有 {placeholder_count} 个占位符")
            return placeholder_count
        except Exception as e:
            self.logger.error(f"读取模板文件时出错：{e}")
            return 11  # 默认值，如果读取模板文件出错

    def run(self):
        try:
            # 创建必要的目录
            download_dir = 'download'
            csvprocess_dir = 'csvprocess'
            os.makedirs(download_dir, exist_ok=True)
            os.makedirs(csvprocess_dir, exist_ok=True)

            # 遍历 download 目录中的所有文件
            for filename in os.listdir(download_dir):
                if filename.endswith('.csv'):
                    file_path = os.path.join(download_dir, filename)
                    
                    try:
                        # 打开 CSV 文件并读取第一行的前两个元素
                        with open(file_path, 'r', encoding='utf-8') as file:
                            csv_reader = csv.reader(file)
                            first_row = next(csv_reader, None)  # 获取第一行

                            # 确保第一行至少有两个元素
                            if first_row and len(first_row) >= 2:
                                # 生成新的文件名
                                new_filename = f"{first_row[0]} {first_row[1]}.csv"
                                # 替换不合法的字符为下划线
                                new_filename = re.sub(r'[\\/*?:"<>|]', "_", new_filename)
                                new_file_path = os.path.join(csvprocess_dir, new_filename)

                                # 复制文件到 csvprocess 目录并重命名
                                shutil.copy2(file_path, new_file_path)
                                self.logger.info(f"文件 {filename} 已复制并重命名为: {new_filename}")
                            else:
                                self.logger.warning(f"文件 {filename} 格式不正确或缺少元素，无法重命名")
                    except Exception as e:
                        self.logger.error(f"处理文件 {filename} 时出错: {e}")

            # 检查处理后的 CSV 文件
            self.check_processed_csvs()

        except Exception as e:
            self.logger.error(f"CSV 处理任务出错: {e}")
        finally:
            self.logger.info("CSV 处理任务结束")

    def check_processed_csvs(self):
        """检查处理后的 CSV 文件列数"""
        csvprocess_dir = 'csvprocess'

        if not os.path.exists(csvprocess_dir):
            self.log_signal.emit(f"警告：文件夹不存在：{csvprocess_dir}")
            return

        for filename in os.listdir(csvprocess_dir):
            if filename.endswith('.csv'):
                file_path = os.path.join(csvprocess_dir, filename)
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        csv_reader = csv.reader(file)
                        first_row = next(csv_reader, None)  # 获取第一行
                        if first_row and len(first_row) != self.expected_columns:
                            self.log_signal.emit(f"警告：文件 '{filename}' 列数不等于模板中的占位符数量 ({self.expected_columns}), 可能有误，请手动检查并更改")
                except Exception as e:
                    self.log_signal.emit(f"检查文件 '{filename}' 时出错：{e}")