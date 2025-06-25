from PySide6.QtCore import QThread, Signal
import os
import csv
import shutil
from docx import Document
from docx.shared import Inches
from utils.log_util import LogUtil

class MakeWordWorker(QThread):
    log_signal = Signal(str)  # 定义日志信号

    def __init__(self):
        super().__init__()
        self.log_util = LogUtil(level='INFO')
        self.logger = self.log_util.get_logger()

    def run(self):
        try:
            # 创建必要的目录
            csvprocess_dir = 'csvprocess'
            output_dir = 'output'
            os.makedirs(csvprocess_dir, exist_ok=True)
            os.makedirs(output_dir, exist_ok=True)

            # Word 模版文件路径
            template_path = 'template.docx'

            # 检查模版文件是否存在
            if not os.path.exists(template_path):
                self.logger.error(f"Word 模版文件 {template_path} 不存在，请确保文件路径正确。")
                return

            # 尝试的编码列表
            encodings_to_try = ['utf-8', 'gbk', 'gb18030']

            # 遍历 csvprocess 目录中的所有 CSV 文件
            for filename in os.listdir(csvprocess_dir):
                if filename.endswith('.csv'):
                    csv_file_path = os.path.join(csvprocess_dir, filename)
                    
                    data = None
                    # 尝试用不同的编码读取文件
                    for encoding in encodings_to_try:
                        try:
                            with open(csv_file_path, 'r', encoding=encoding) as csvfile:
                                csvreader = csv.reader(csvfile)
                                data = list(csvreader)
                            self.logger.info(f"成功使用编码 {encoding} 读取文件: {filename}")
                            break
                        except UnicodeDecodeError:
                            continue
                    
                    if data is None:
                        self.logger.error(f"无法读取文件 {filename}，跳过处理。")
                        continue
                    
                    try:
                        # 创建一个新的 Word 文档（基于模版）
                        doc = Document(template_path)
                        
                        # 遍历文档中的每个段落，替换占位符
                        for i, paragraph in enumerate(doc.paragraphs):
                            for j, cell in enumerate(data[0]):  # 假设使用第一行数据
                                placeholder = f"{{{{{j+1}}}}}"
                                if placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, cell)
                        
                        # 遍历文档中的每个表格，替换占位符
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        for j, cell_data in enumerate(data[0]):
                                            placeholder = f"{{{{{j+1}}}}}"
                                            if placeholder in paragraph.text:
                                                paragraph.text = paragraph.text.replace(placeholder, cell_data)
                        
                        # 生成新的 Word 文件名
                        new_word_filename = f"{os.path.splitext(filename)[0]}.docx"
                        new_word_path = os.path.join(output_dir, new_word_filename)
                        
                        # 保存新的 Word 文档
                        doc.save(new_word_path)
                        self.logger.info(f"生成新的 Word 文档: {new_word_filename}")
                    
                    except Exception as e:
                        self.logger.error(f"处理文件 {filename} 时出错: {e}")
        except Exception as e:
            self.logger.error(f"Word 生成任务出错: {e}")
        finally:
            self.logger.info("Word 生成任务结束")